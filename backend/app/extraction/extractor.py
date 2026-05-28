import io
import json
import re
import csv

from groq import Groq

from app.core.config import get_settings
from app.models import ExtractedInvoiceData, LineItemExtracted


# temperature=0.1 keeps extraction deterministic across runs.
EXTRACTION_PROMPT = """You are an invoice data extraction engine. Extract structured data from the invoice text below.

Return ONLY a valid JSON object with this exact schema:
{
  "invoice_number": "string or null",
  "vendor_name": "string or null",
  "vendor_email": "string or null",
  "invoice_date": "YYYY-MM-DD string or null",
  "po_reference": "string or null",
  "line_items": [
    {
      "sequence_number": 1,
      "item_category": "steel_rod | steel_plate | cement | other",
      "description": "full item description",
      "quantity": numeric_value_or_null,
      "quantity_unit": "MT | KG | PCS | M | M2 | null",
      "unit_rate": numeric_value_or_null,
      "rate_unit": "INR/MT | INR/KG | INR/PCS | null",
      "total_value": numeric_value_or_null,
      "dimensions": {
        "length": numeric_or_null,
        "width": numeric_or_null,
        "thickness": numeric_or_null,
        "diameter": numeric_or_null,
        "unit": "mm | cm | m | inch | null"
      },
      "quality_grade": "grade string or null"
    }
  ],
  "extraction_confidence": "high | medium | low"
}

Rules:
- All numeric values must be actual numbers, not strings
- If a field is not found in the document, use null
- item_category should be your best inference from the description
- extraction_confidence: "high" if most fields found, "medium" if partial, "low" if minimal data

Invoice text:
"""

SUPPORTED_TYPES = {
    "application/pdf":                                                   "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword":                                                "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel":                                          "xlsx",
    "application/json":                                                  "json",
    "text/plain":                                                        "txt",
    "text/csv":                                                          "csv",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".json", ".txt", ".csv"}


# pymupdf (fitz) is used as the primary PDF parser — ships as a self-contained wheel,
# no system binaries needed. For image-only pages, falls back to pytesseract if available.
def _pdf_to_text(file_bytes: bytes) -> str:
    import fitz  # pymupdf

    try:
        import pytesseract
        from PIL import Image as PILImage
        import io as _io
        ocr_available = True
    except (ImportError, Exception):
        ocr_available = False

    text_parts = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    for page in doc:
        page_text = page.get_text("text").strip()

        # b = (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")
        block_lines = []
        for b in blocks:
            if b[6] == 0 and b[4].strip():  # type 0 = text block
                block_lines.append(b[4].strip().replace("\n", " "))

        if page_text:
            text_parts.append(page_text)
        elif block_lines:
            text_parts.extend(block_lines)
        elif ocr_available:
            # Truly image-only page — render at 200 DPI and run tesseract
            mat      = fitz.Matrix(200 / 72, 200 / 72)
            pix      = page.get_pixmap(matrix=mat)
            img      = PILImage.open(_io.BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(img, lang="eng")
            if ocr_text.strip():
                text_parts.append(f"[OCR] {ocr_text.strip()}")
        else:
            text_parts.append("[image-only page — install tesseract for OCR]")

    doc.close()
    return "\n".join(text_parts)


# Tables rendered as pipe-delimited rows for consistent LLM parsing.
def _docx_to_text(file_bytes: bytes) -> str:
    from docx import Document  # python-docx

    doc   = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# openpyxl covers .xlsx; xlrd covers legacy .xls.
def _xlsx_to_text(file_bytes: bytes) -> str:
    import openpyxl

    wb    = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    parts = []

    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _json_to_text(file_bytes: bytes) -> str:
    try:
        data = json.loads(file_bytes.decode("utf-8", errors="replace"))
        return json.dumps(data, indent=2)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON file: {e}")


def _txt_to_text(file_bytes: bytes, filename: str) -> str:
    raw = file_bytes.decode("utf-8", errors="replace")

    if filename.lower().endswith(".csv"):
        lines  = raw.splitlines()
        reader = csv.reader(lines)
        return "\n".join(" | ".join(row) for row in reader if any(row))

    return raw


# Determines format from filename extension (more reliable than MIME on uploads).
def _file_to_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()

    if name.endswith(".pdf"):
        return _pdf_to_text(file_bytes)
    if name.endswith((".docx", ".doc")):
        return _docx_to_text(file_bytes)
    if name.endswith((".xlsx", ".xls")):
        return _xlsx_to_text(file_bytes)
    if name.endswith(".json"):
        return _json_to_text(file_bytes)
    if name.endswith((".txt", ".csv")):
        return _txt_to_text(file_bytes, filename)

    raise ValueError(f"Unsupported file type: {filename}")


# Groq may wrap JSON in ```json fences or prefix with chatter — strip fences
# first, then find the outermost balanced {...} block.
def _parse_json(content: str) -> dict:
    text = content.strip()

    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fence:
        return json.loads(fence.group(1))

    start = text.find("{")
    if start == -1:
        raise ValueError("Groq response contained no JSON object")
    depth = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return json.loads(text[start : i + 1])
    raise ValueError("Groq response had unbalanced JSON braces")


# filename is passed so we can route to the correct parser.
def extract_invoice_data(file_bytes: bytes, filename: str = "invoice.pdf") -> ExtractedInvoiceData:
    settings = get_settings()
    client   = Groq(api_key=settings.groq_api_key)

    raw_text = _file_to_text(file_bytes, filename)

    # Truncate to 6000 chars to stay within Groq context limits.
    prompt = EXTRACTION_PROMPT + raw_text[:6000]

    # Retry on rate limit — free tier is 30 RPM.
    import time
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=settings.groq_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=2048,
            )
            break
        except Exception as e:
            if attempt < 2 and ("rate" in str(e).lower() or "temporarily" in str(e).lower() or "429" in str(e)):
                time.sleep(8)
                continue
            raise

    raw_json = _parse_json(response.choices[0].message.content)

    line_items = []
    for i, item in enumerate(raw_json.get("line_items", [])):
        dims = item.get("dimensions") or {}
        line_items.append(LineItemExtracted(
            sequence_number=item.get("sequence_number", i + 1),
            item_category=item.get("item_category"),
            description=item.get("description"),
            quantity=item.get("quantity"),
            quantity_unit=item.get("quantity_unit"),
            unit_rate=item.get("unit_rate"),
            rate_unit=item.get("rate_unit"),
            total_value=item.get("total_value"),
            # Only store dimensions if at least one value is non-null
            dimensions=dims if any(dims.values()) else None,
            quality_grade=item.get("quality_grade"),
        ))

    return ExtractedInvoiceData(
        invoice_number=raw_json.get("invoice_number") or "UNKNOWN",
        vendor_name=raw_json.get("vendor_name") or "Unknown Vendor",
        vendor_email=raw_json.get("vendor_email"),
        invoice_date=raw_json.get("invoice_date"),
        po_reference=raw_json.get("po_reference"),
        line_items=line_items,
        extraction_confidence=raw_json.get("extraction_confidence", "medium"),
        raw_text_length=len(raw_text),
    )
